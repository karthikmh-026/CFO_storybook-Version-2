import os
import openpyxl
import csv
import docx  # python-docx might not be installed, we will handle ImportError
import sys

search_dir = r"C:\Users\ajala\Downloads"
search_terms = ["1965.44", "1856.82", "320.19", "3201933", "3,201,933", "1965", "1856"]
excludes = ['node_modules', 'build', '.git', '.conda', 'anaconda3', 'Vittora', 'SAPGUI750']

def search_text(content, filepath, item_name):
    for term in search_terms:
        if term in content:
            print(f"Match in {filepath} ({item_name}): found term '{term}'")

print("Starting deep search in Downloads...")

for root, dirs, files in os.walk(search_dir):
    for exc in excludes:
        if exc in dirs:
            dirs.remove(exc)
            
    for file in files:
        filepath = os.path.join(root, file)
        
        # Skip very large files
        try:
            sz = os.path.getsize(filepath)
            if sz > 50 * 1024 * 1024:  # 50MB
                continue
        except:
            continue
            
        ext = file.lower()
        
        if ext in ('.txt', '.csv', '.json', '.xml', '.py', '.js', '.jsx', '.html', '.css'):
            try:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                    search_text(content, filepath, "text content")
            except Exception as e:
                pass
                
        elif ext in ('.xlsx', '.xlsm'):
            try:
                wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    for r_idx, row in enumerate(ws.iter_rows(values_only=True), 1):
                        for c_idx, val in enumerate(row, 1):
                            if val is not None:
                                val_str = str(val)
                                for term in search_terms:
                                    if term in val_str:
                                        print(f"Match in Excel: {filepath} | Sheet: {sheet} | Cell ({r_idx},{c_idx}) = {val}")
            except Exception as e:
                pass
                
        elif ext == '.docx':
            try:
                doc = docx.Document(filepath)
                fullText = []
                for para in doc.paragraphs:
                    fullText.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            fullText.append(cell.text)
                content = "\n".join(fullText)
                search_text(content, filepath, "Word document")
            except ImportError:
                pass
            except Exception as e:
                pass
                
        elif ext == '.pdf':
            # Try parsing using standard pypdf if installed
            try:
                import pypdf
                reader = pypdf.PdfReader(filepath)
                content = ""
                for page in reader.pages:
                    content += page.extract_text() or ""
                search_text(content, filepath, "PDF document")
            except:
                pass

print("Search complete.")
